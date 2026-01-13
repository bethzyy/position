"""
简历解析模块 - 支持PDF、Word、TXT格式
"""
import os
import logging
from pathlib import Path
from typing import Optional


class ResumeParser:
    """简历解析器"""

    def __init__(self):
        """初始化简历解析器"""
        self.logger = logging.getLogger(__name__)

    def parse_resume(self, file_path: str) -> str:
        """
        解析简历文件

        Args:
            file_path: 简历文件路径

        Returns:
            简历文本内容
        """
        self.logger.info(f"开始解析简历: {file_path}")

        file_path_obj = Path(file_path)

        if not file_path_obj.exists():
            raise FileNotFoundError(f"简历文件不存在: {file_path}")

        # 根据文件扩展名选择解析方法
        file_ext = file_path_obj.suffix.lower()

        try:
            if file_ext == '.pdf':
                content = self._parse_pdf(file_path)
            elif file_ext in ['.doc', '.docx']:
                content = self._parse_word(file_path)
            elif file_ext in ['.txt', '.text']:
                content = self._parse_text(file_path)
            else:
                raise ValueError(f"不支持的文件格式: {file_ext}")

            self.logger.info(f"简历解析成功，内容长度: {len(content)} 字符")
            return content

        except Exception as e:
            self.logger.error(f"简历解析失败: {e}", exc_info=True)
            raise

    def _parse_pdf(self, file_path: str) -> str:
        """
        解析PDF简历

        Args:
            file_path: PDF文件路径

        Returns:
            文本内容
        """
        try:
            import PyPDF2
        except ImportError:
            raise ImportError(
                "未安装PyPDF2库，请运行: pip install PyPDF2"
            )

        text_content = []

        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)

                # 获取总页数
                num_pages = len(pdf_reader.pages)
                self.logger.info(f"PDF总页数: {num_pages}")

                # 提取每一页的文本
                for page_num in range(num_pages):
                    page = pdf_reader.pages[page_num]
                    text = page.extract_text()
                    if text.strip():
                        text_content.append(text)

        except Exception as e:
            self.logger.error(f"PDF解析出错: {e}")
            raise

        return '\n'.join(text_content)

    def _parse_word(self, file_path: str) -> str:
        """
        解析Word简历

        Args:
            file_path: Word文件路径

        Returns:
            文本内容
        """
        try:
            from docx import Document
        except ImportError:
            raise ImportError(
                "未安装python-docx库，请运行: pip install python-docx"
            )

        try:
            doc = Document(file_path)
            text_content = []

            # 提取段落文本
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)

            # 提取表格文本
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells]
                    text_content.append(' | '.join(row_text))

        except Exception as e:
            self.logger.error(f"Word文档解析出错: {e}")
            raise

        return '\n'.join(text_content)

    def _parse_text(self, file_path: str) -> str:
        """
        解析文本简历

        Args:
            file_path: 文本文件路径

        Returns:
            文本内容
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
        except UnicodeDecodeError:
            # 尝试其他编码
            try:
                with open(file_path, 'r', encoding='gbk') as f:
                    content = f.read()
            except:
                with open(file_path, 'r', encoding='latin-1') as f:
                    content = f.read()

        return content


def test_resume_parser():
    """测试函数"""
    logging.basicConfig(
        level=logging.INFO,
        format='%(asctime)s - %(levelname)s - %(message)s'
    )

    parser = ResumeParser()

    # 创建测试文本简历
    test_resume = Path("tests/test_resume.txt")
    test_resume.parent.mkdir(exist_ok=True)

    with open(test_resume, 'w', encoding='utf-8') as f:
        f.write("""
张三
高级软件工程师

教育背景:
XX大学 计算机科学与技术 本科 2015-2019

工作经历:
ABC公司 软件工程师 2019-2022
- 负责后端开发
- 使用Python和Django框架

XYZ公司 高级软件工程师 2022-至今
- 负责系统架构设计
- 使用微服务架构

技能:
- Python, Java, Go
- MySQL, Redis, MongoDB
- Docker, Kubernetes
        """)

    try:
        content = parser.parse_resume(str(test_resume))
        print("简历内容:")
        print(content)
    except Exception as e:
        print(f"测试失败: {e}")


if __name__ == "__main__":
    test_resume_parser()
